import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import baza

SLIKE_POT = Path(__file__).parent.parent / "slike"

PODPRTI_FORMATI = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"}


def _preveri_slike(besedilo: str, slike_po_imenu: dict) -> list[str]:
    """Vrne seznam manjkajočih ali nepodprtih slik za nalogo."""
    napake = []
    for m in re.finditer(r'\[SLIKA:([^\]]+)\]', besedilo):
        ime = m.group(1).strip()
        pot = slike_po_imenu.get(ime)
        if not pot or not pot.exists():
            napake.append(f"slika ne obstaja: {ime}")
        elif pot.suffix.lower() not in PODPRTI_FORMATI:
            napake.append(f"format ni podprt ({pot.suffix}): {ime}")
    return napake


# ---------------------------------------------------------------------------
# Markdown tabele
# ---------------------------------------------------------------------------

def _je_tabela_vrstica(vrstica: str) -> bool:
    s = vrstica.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def _je_separator(vrstica: str) -> bool:
    return bool(re.match(r'^\|[\s\-:]+(\|[\s\-:]+)*\|$', vrstica.strip()))


def _razcleni_md_vrstico(vrstica: str) -> list[str]:
    v = vrstica.strip()
    if v.startswith("|"):
        v = v[1:]
    if v.endswith("|"):
        v = v[:-1]
    return [c.strip() for c in v.split("|")]


def _razdeli_na_bloke(tekst: str) -> list[tuple[str, object]]:
    """Razdeli besedilo na bloke: ('text', str) ali ('tabela', list[str])."""
    vrstice = tekst.split("\n")
    bloki = []
    i = 0
    while i < len(vrstice):
        if _je_tabela_vrstica(vrstice[i]):
            tabela_vrstice = []
            while i < len(vrstice) and _je_tabela_vrstica(vrstice[i]):
                tabela_vrstice.append(vrstice[i])
                i += 1
            bloki.append(("tabela", tabela_vrstice))
        else:
            tekst_vrstice = []
            while i < len(vrstice) and not _je_tabela_vrstica(vrstice[i]):
                tekst_vrstice.append(vrstice[i])
                i += 1
            skupaj = "\n".join(tekst_vrstice).strip()
            if skupaj:
                bloki.append(("text", skupaj))
    return bloki


def _dodaj_md_tabelo(doc: Document, md_vrstice: list[str]):
    """Pretvori markdown vrstice v python-docx tabelo."""
    podatki = [v for v in md_vrstice if not _je_separator(v)]
    if not podatki:
        return

    vrstice_celic = [_razcleni_md_vrstico(v) for v in podatki]
    st_stolpcev = max(len(v) for v in vrstice_celic)
    if st_stolpcev == 0:
        return

    tabela = doc.add_table(rows=len(vrstice_celic), cols=st_stolpcev)
    try:
        tabela.style = "Table Grid"
    except Exception:
        pass

    for i, vrstica in enumerate(vrstice_celic):
        for j in range(st_stolpcev):
            vrednost = vrstica[j] if j < len(vrstica) else ""
            celica = tabela.cell(i, j)
            celica.text = vrednost
            if i == 0 and vrednost:
                for run in celica.paragraphs[0].runs:
                    run.bold = True


# ---------------------------------------------------------------------------
# Vstavljanje besedila in slik
# ---------------------------------------------------------------------------

def _dodaj_besedilo_s_slikami(doc: Document, besedilo: str, slike_po_imenu: dict, stevilka: int):
    """Razbije besedilo na segmente (slike, tabele, tekst) in jih vstavi v dokument."""
    segmenti = re.split(r'(\[SLIKA:[^\]]+\])', besedilo)

    prvi = True  # Številka naloge še ni bila dodana

    for segment in segmenti:
        m = re.match(r'\[SLIKA:([^\]]+)\]', segment)
        if m:
            # Slika — pred njo dodaj številko če je prva stvar v nalogi
            if prvi:
                odst = doc.add_paragraph()
                odst.paragraph_format.space_before = Pt(6)
                odst.add_run(f"{stevilka}. ").bold = True
                prvi = False
            ime = m.group(1).strip()
            pot = slike_po_imenu.get(ime)
            if pot and pot.exists() and pot.suffix.lower() in PODPRTI_FORMATI:
                doc.add_picture(str(pot), width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            tekst = segment.strip()
            if not tekst:
                continue
            bloki = _razdeli_na_bloke(tekst)
            for tip, vsebina in bloki:
                if tip == "tabela":
                    if prvi:
                        odst = doc.add_paragraph()
                        odst.paragraph_format.space_before = Pt(6)
                        odst.add_run(f"{stevilka}. ").bold = True
                        prvi = False
                    _dodaj_md_tabelo(doc, vsebina)
                else:
                    if prvi:
                        odst = doc.add_paragraph()
                        odst.paragraph_format.space_before = Pt(6)
                        odst.add_run(f"{stevilka}. ").bold = True
                        odst.add_run(vsebina)
                        prvi = False
                    else:
                        doc.add_paragraph(vsebina)

    if prvi:
        odst = doc.add_paragraph()
        odst.paragraph_format.space_before = Pt(6)
        odst.add_run(f"{stevilka}. ").bold = True


def generiraj_test(ids_nalog: list[int], naslov: str = "Test iz biologije") -> tuple[bytes, list[str]]:
    """Ustvari .docx test z izbranimi nalogami.

    Vrne (vsebina_docx, seznam_napak).
    Naloge z manjkajočimi/nepodprtimi slikami so izpuščene iz dokumenta.
    """
    naloge = baza.pridobi_naloge_po_ids(ids_nalog)

    doc = Document()
    naslov_odst = doc.add_heading(naslov, level=0)
    naslov_odst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    napake = []
    stevilka = 1
    for naloga in naloge:
        slike_po_imenu = {}
        if naloga["ima_sliko"]:
            for slika in baza.pridobi_slike_naloge(naloga["id"]):
                pot = SLIKE_POT / slika["ime_datoteke"]
                slike_po_imenu[pot.name] = pot

        napake_naloge = _preveri_slike(naloga["besedilo"], slike_po_imenu)
        if napake_naloge:
            napake.append(f"Naloga {naloga['id']}: {'; '.join(napake_naloge)} — izpuščena")
            continue

        _dodaj_besedilo_s_slikami(doc, naloga["besedilo"], slike_po_imenu, stevilka)
        stevilka += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read(), napake
